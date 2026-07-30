import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement

from core.blog import article_toc_and_anchors
from core.page_blueprint import build_page_blueprint
from ingestion.blog import parse_blog_document, parse_blog_docx
from renderer.engine import render_resolved


class BlogPipelineTests(unittest.TestCase):
    def test_structural_parser_preserves_article_content_without_editor_metadata(self):
        payload = parse_blog_document([
            {"type": "h1", "text": "A structured article"},
            {"type": "paragraph", "text": "A concise introduction from the source document."},
            {"type": "h2", "text": "A factual section"},
            {"type": "paragraph", "text": "The article body remains source-authored."},
            {"type": "table", "table_title": "", "headers": ["Name", "Value"], "rows": [["One", "1"]]},
            {"type": "h2", "text": "Reader questions"},
            {"type": "list_item", "text": "What is the first question?", "list_kind": "ul", "list_level": 0},
            {"type": "list_item", "text": "The first source-backed answer contains enough useful detail.", "list_kind": "ul", "list_level": 0},
            {"type": "list_item", "text": "What is the second question?", "list_kind": "ul", "list_level": 0},
            {"type": "list_item", "text": "The second source-backed answer also contains useful detail.", "list_kind": "ul", "list_level": 0},
        ], "article.docx")

        self.assertEqual(payload["title"], "A structured article")
        self.assertEqual(payload["subtitle"], "A concise introduction from the source document.")
        self.assertEqual(len(payload["faqs"]), 2)
        self.assertIn("<table><thead>", payload["content_html"])
        self.assertIn("<th scope=\"col\">Name</th>", payload["content_html"])
        self.assertNotIn("What is the first question?", payload["content_html"])
        self.assertGreater(payload["word_count"], 0)
        self.assertFalse({"author", "published_date", "category", "seo_title", "meta_description", "read_time"} & set(payload))

    def test_docx_reader_does_not_apply_course_table_header_normalisation(self):
        document = Document()
        document.add_paragraph("The article title", style="Title")
        table = document.add_table(rows=2, cols=2)
        table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
        table.rows[0].cells[0].text = "Course Name"
        table.rows[0].cells[1].text = "Fee"
        table.rows[1].cells[0].text = "Online MBA"
        table.rows[1].cells[1].text = "INR 100"

        with tempfile.TemporaryDirectory() as directory:
            filepath = Path(directory) / "article.docx"
            document.save(filepath)
            blocks = parse_blog_docx(str(filepath))

        table_block = next(block for block in blocks if block["type"] == "table")
        self.assertEqual(table_block["headers"], ["Course Name", "Fee"])
        self.assertEqual(table_block["rows"], [["Online MBA", "INR 100"]])

    def test_toc_and_rendering_use_derived_article_data_and_resolved_links(self):
        content_html, toc = article_toc_and_anchors(
            "<h2>First section</h2><p>Article body.</p><h3>Detail</h3><p>More detail.</p>"
        )
        self.assertEqual([(entry["text"], entry["level"]) for entry in toc], [("First section", 2), ("Detail", 3)])
        self.assertIn('id="first-section"', content_html)

        html = render_resolved({
            "slug": "structured-article",
            "page_type": "blog",
            "university_slug": "test",
            "parent_slug": None,
            "raw": {
                "title": "Structured article",
                "content_html": content_html,
                "hero_image_url": "/assets/images/blog-structured.webp",
                "related_course_slugs": ["online-mba"],
                "_workspace_courses": [{
                    "slug": "online-mba",
                    "data": {
                        "program_name": "Online MBA",
                        "hero_description": "A real course description.",
                        "duration": "2 years",
                        "mode": "Online",
                        "total_fee": "INR 100",
                    },
                }],
            },
        }, standalone=True)

        self.assertIn("Online MBA", html)
        self.assertIn('id="first-section"', html)
        self.assertIn('"@type":"Article"', html)
        self.assertNotIn("Krishna Porwal", html)

    def test_blueprint_owns_every_blog_template_field(self):
        blueprint = build_page_blueprint("blog")
        self.assertEqual(blueprint["fields"]["title"]["source"], "AUTO")
        self.assertEqual(blueprint["fields"]["author"]["source"], "MANUAL")
        self.assertEqual(blueprint["external_fields"]["toc"]["source"], "DERIVED")
        self.assertEqual(blueprint["external_fields"]["related_courses"]["source"], "WORKSPACE")


if __name__ == "__main__":
    unittest.main()
