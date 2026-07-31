import asyncio
from io import BytesIO
import unittest
from unittest.mock import patch

from docx import Document
from fastapi import HTTPException, UploadFile

from ingestion.extractor import extract_acf
from main import IngestRequest, ingest_acf, parse_docx_endpoint, validate_blueprint_content
from renderer.engine import clean_spec_name, derive_financing_from_emi, parse_admission_html


class PostParserPipelineTests(unittest.TestCase):
    def test_endpoint_validation_reports_all_blueprint_required_fields(self):
        with self.assertRaises(HTTPException) as context:
            validate_blueprint_content(
                "course",
                {"program_name": "Online MBA"},
                slug="ignou-online-mba",
                university_slug="ignou",
                parent_slug=None,
            )

        self.assertEqual(context.exception.status_code, 422)
        missing = {
            entry["field"] for entry in context.exception.detail["missing_fields"]
        }
        self.assertIn("hero_description", missing)
        self.assertIn("duration", missing)
        self.assertIn("total_fee", missing)
        self.assertIn("hero_image_url", missing)
        self.assertIn("certificate_image_url", missing)

    def test_blog_validation_requires_title_body_and_featured_image(self):
        with self.assertRaises(HTTPException) as context:
            validate_blueprint_content(
                "blog",
                {"title": "An article"},
                slug="an-article",
                university_slug="ignou",
                parent_slug=None,
            )

        self.assertEqual(context.exception.status_code, 422)
        self.assertEqual(
            {entry["field"] for entry in context.exception.detail["missing_fields"]},
            {"content_html", "hero_image_url"},
        )

    def test_ingest_response_carries_field_state_without_changing_acf_data(self):
        payload = {
            "page_type": "university",
            "slug": "ignou",
            "university_slug": "ignou",
            "university_name": "IGNOU",
            "hero_description": "A public open university.",
        }

        result = asyncio.run(ingest_acf(IngestRequest(acf_data=payload)))

        self.assertEqual(result["status"], "ok")
        self.assertEqual(result["acf_data"]["hero_description"], payload["hero_description"])
        self.assertEqual(result["field_state"]["slug"]["value"], "ignou")
        self.assertTrue(result["field_state"]["hero_image_url"]["missing"])
        self.assertTrue(result["field_state"]["hero_image_url"]["manual"])

        # Phase 2: page_state rides alongside field_state without touching acf_data.
        page_state = result["page_state"]
        self.assertEqual(page_state["page_type"], "university")
        hero = next(s for s in page_state["sections"] if s["section"] == "hero")
        self.assertTrue(hero["renderable"])  # name + description present
        # hero_image_url is required by the builder but optional to the template.
        self.assertIn("hero_image_url", hero["missing_optional"])

    def test_course_docx_response_builds_state_after_parser_result(self):
        document = Document()
        document.add_heading("Online MBA", level=1)
        document.add_paragraph("A flexible management programme.")
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        upload = UploadFile(filename="online-mba.docx", file=buffer)
        micro_result = {
            "filename": "online-mba.docx",
            "payload": {
                "program_name": "Online MBA",
                "university_name": "IGNOU",
                "hero_description": "A flexible management programme.",
                "reviews": [],
            },
        }

        with patch("main.forward_to_micro_pipeline", return_value=micro_result):
            result = asyncio.run(
                parse_docx_endpoint(
                    file=upload,
                    page_type="course",
                    university_slug="ignou",
                )
            )

        self.assertEqual(result["payload"]["program_name"], "Online MBA")
        self.assertEqual(result["field_state"]["university_slug"]["value"], "ignou")
        self.assertEqual(result["field_state"]["slug"]["value"], "ignou-online-mba")
        self.assertEqual(result["field_state"]["reviews"]["source"], "MANUAL")
        self.assertEqual(result["table_warnings"], [])

    def test_new_micro_payload_is_normalized_before_review_state(self):
        document = Document()
        document.add_heading("Online MBA", level=1)
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        upload = UploadFile(filename="online-mba.docx", file=buffer)
        micro_result = {
            "filename": "online-mba.docx",
            "payload": {
                "program_name": "Online MBA",
                "university_name": "IGNOU",
                "hero_image": "https://cdn.example.test/new-hero.webp",
                "hero_image_url": "https://cdn.example.test/legacy-hero.webp",
                "certificate_image": "https://cdn.example.test/certificate.webp",
                "ugc_approved": "UGC Entitled",
                "mode_of_learning": "Online",
                "eligibility_content": [
                    {
                        "eligibility_title": "Academic requirement",
                        "eligibility_description": "A recognised bachelor's degree.",
                    }
                ],
                "fee_plans": [
                    {
                        "plan_name": "Semester I",
                        "plan_amount": "INR 10,000",
                    }
                ],
                "reviews": [
                    {
                        "review_text": "Helpful course structure.",
                        "reviewer_name": "Rahul Verma",
                        "reviewer_label": "MBA Student",
                    }
                ],
            },
        }

        with patch("main.forward_to_micro_pipeline", return_value=micro_result):
            result = asyncio.run(
                parse_docx_endpoint(
                    file=upload,
                    page_type="course",
                    university_slug="ignou",
                )
            )

        payload = result["payload"]
        self.assertEqual(payload["hero_image_url"], "https://cdn.example.test/new-hero.webp")
        self.assertEqual(payload["certificate_image_url"], "https://cdn.example.test/certificate.webp")
        self.assertEqual(payload["ugc_status"], "UGC Entitled")
        self.assertEqual(payload["mode"], "Online")
        self.assertNotIn("hero_image", payload)
        self.assertNotIn("ugc_approved", payload)
        self.assertNotIn("mode_of_learning", payload)
        self.assertEqual(payload["eligibility_content"][0]["eligibility_title"], "Academic requirement")
        self.assertEqual(result["field_state"]["fee_plans"]["value"][0]["plan_name"], "Semester I")
        self.assertEqual(payload["reviews"][0]["reviewer_name"], "Rahul Verma")

    def test_singular_fee_tag_attaches_its_table_to_course_fee_plans(self):
        acf = extract_acf([
            {"type": "h2", "text": "[fee_heading,fee_plans] Fee Structure"},
            {
                "type": "table",
                "headers": ["Semester", "Fee"],
                "rows": [
                    ["Semester I", "INR 16,000/-"],
                    ["Semester II", "INR 16,000/-"],
                ],
            },
        ], "course", {})

        self.assertEqual(acf["fee_plans"], [
            {"plan_name": "Semester I", "plan_amount": "INR 16,000/-", "plan_total": "INR 16,000/-"},
            {"plan_name": "Semester II", "plan_amount": "INR 16,000/-", "plan_total": "INR 16,000/-"},
        ])

    def test_explicit_identity_marker_overrides_conflicting_micro_value(self):
        document = Document()
        document.add_heading("[university_name] LPU Online", level=1)
        document.add_paragraph("University Full Name: Lovely Professional University Online")
        document.add_paragraph("NAAC Grade: A++")
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        upload = UploadFile(filename="lpu.docx", file=buffer)
        micro_result = {
            "filename": "lpu.docx",
            "payload": {
                "university_name": "Lovely Professional Online",
                "hero_description": "Online learning programmes.",
                "naac_grade": "A++",
                "ugc_approved": "Entitled",
            },
        }

        with patch("main.forward_to_micro_pipeline", return_value=micro_result):
            result = asyncio.run(
                parse_docx_endpoint(
                    file=upload,
                    page_type="university",
                    university_slug="lpu",
                )
            )

        self.assertEqual(result["payload"]["university_name"], "LPU Online")

    def test_admission_note_sections_do_not_become_numbered_steps(self):
        steps = parse_admission_html(
            "<p>Step 1. Register online.</p>"
            "<p>Step 2. Submit the application.</p>"
            "<h4>Important Notes:</h4>"
            "<ul><li>The registration fee is non-refundable.</li></ul>"
        )

        self.assertEqual([step["t"] for step in steps], [
            "Register online.",
            "Submit the application.",
        ])

    def test_specialization_display_name_omits_parenthetical_parser_suffixes(self):
        self.assertEqual(clean_spec_name("Finance (Fin"), "Finance")
        self.assertEqual(clean_spec_name("Information Technology (IT)"), "Information Technology")

    def test_emi_content_derives_only_explicit_financing_facts(self):
        financing, banks = derive_financing_from_emi(
            "<p>EMI per month ranges between INR 7373 /-. EMI can be opted for "
            "a period of 12 to 24 months. The university has bank recommendations:</p>"
            "<ul><li>Punjab National Bank</li><li>IDBI Bank</li></ul>"
            "<p>The university offers up to 30% of student grants and scholarships.</p>"
        )

        self.assertEqual([item["stat"] for item in financing], [
            "₹7,373", "12–24 months", "Up to 30%",
        ])
        self.assertEqual(banks, ["Punjab National Bank", "IDBI Bank"])


if __name__ == "__main__":
    unittest.main()
